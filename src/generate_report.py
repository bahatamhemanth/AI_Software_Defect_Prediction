"""Generate the final technical report from persisted measured artifacts."""

import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))
from data_preprocessing import load_raw_dataset  # noqa: E402


def add_table(document, frame):
    table = document.add_table(rows=1, cols=len(frame.columns))
    table.style = "Light Shading Accent 1"
    for cell, name in zip(table.rows[0].cells, frame.columns):
        cell.text = str(name)
    for _, row in frame.iterrows():
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = f"{value:.4f}" if isinstance(value, float) else str(value)


def main() -> None:
    import pandas as pd

    reports = ROOT / "reports"
    figures = reports / "figures"
    data = load_raw_dataset()
    results = pd.read_csv(reports / "model_results.csv")
    importance = pd.read_csv(reports / "feature_importance.csv").head(10)
    summary = json.loads((reports / "preprocessing_summary.json").read_text(encoding="utf-8"))
    best = results.iloc[0]
    document = Document()
    document.add_heading("AI-Based Software Defect Prediction and Test Prioritization System", 0)
    document.add_paragraph("Technical Implementation and Evaluation Report")
    document.add_heading("Abstract", 1)
    document.add_paragraph("This project develops a reproducible machine-learning workflow for identifying software modules at risk of defects and prioritizing testing. The implementation uses the public NASA MDP PC1 dataset, leakage-safe preprocessing, five classification models, recall-aware model selection, probability-based risk categories, and a Streamlit demonstration.")
    document.add_heading("Chapter 1: Introduction", 1)
    document.add_paragraph("Software defect prediction uses historical software metrics to identify modules that deserve additional quality assurance attention. In this project, the objective is binary classification of PC1 modules as defective or non-defective, followed by risk-based test prioritization. Since defective modules are a minority, missing a defective module is treated as an important error.")
    document.add_heading("Chapter 2: Dataset and Methodology", 1)
    document.add_paragraph(f"NASA MDP PC1 contains {len(data):,} modules, 21 numeric software-metric predictors, and the binary defects target. The measured class distribution is {int((data.defects == False).sum()):,} non-defective modules and {int((data.defects == True).sum()):,} defective modules ({data.defects.mean() * 100:.2f}% defective). The raw ARFF file is retained unchanged under data/raw/pc1.arff.")
    document.add_paragraph("EDA found 0 missing values, 155 duplicate rows, 0 infinite values, pronounced right skew in several metrics, and perfect correlation between E and T. Duplicate rows were removed for modeling, T was excluded as exact redundancy, plausible extreme values were retained, and median imputation plus standard scaling were placed inside the fitted pipeline. Supported estimators used balanced class weights; no resampling was applied.")
    document.add_paragraph("The data was split stratifiably with random_state=42. Five-fold stratified cross-validation on the training partition selected the model by defective-class recall, then average precision and F1. The held-out test partition was used once for final metrics.")
    document.add_heading("Chapter 3: Machine Learning Implementation", 1)
    document.add_paragraph("The compared classifiers were Logistic Regression, Decision Tree, Random Forest, Gradient Boosting, and Support Vector Machine. Each estimator was evaluated through the same preprocessing pipeline and generated defect probabilities where supported.")
    document.add_heading("Chapter 4: Results and Discussion", 1)
    document.add_paragraph(f"The selected model was {best['model']}. Its held-out test results were accuracy {best['accuracy']:.4f}, precision {best['precision']:.4f}, recall {best['recall']:.4f}, F1 {best['f1']:.4f}, ROC-AUC {best['roc_auc']:.4f}, and PR-AUC {best['pr_auc']:.4f}. Recall-first selection was used because the defective class is rare and missed defects are costly.")
    add_table(document, results)
    for filename, caption in [("target_distribution.png", "PC1 target distribution"), ("feature_distributions.png", "PC1 metric distributions"), ("correlation_heatmap.png", "Metric correlation heatmap"), ("model_curves.png", "Selected model ROC and precision-recall curves"), ("feature_importance.png", "Selected model feature importance")]:
        path = figures / filename
        if path.exists():
            document.add_picture(str(path), width=Inches(6.2))
            document.add_paragraph(caption)
    document.add_paragraph("The strongest observed feature correlation was E with T (r=1.000). The most skewed metrics included T, E, ev(g), iv(G), and V. IQR analysis flagged extreme observations, but these were retained because an extreme software metric may be valid rather than erroneous.")
    document.add_heading("Chapter 5: Prediction and Test Prioritization", 1)
    document.add_paragraph("The persisted joblib bundle accepts the trained PC1 metric columns and returns a predicted class and defect probability. Risk levels are Low below 0.33, Medium from 0.33 to below 0.66, and High at or above 0.66. These thresholds are transparent operational bands, not claims of calibrated probability boundaries.")
    document.add_paragraph("Testing priority uses defect probability as the primary signal and combines it with a normalized complexity index calculated from loc, v(g), N, V, and branchCount. The implemented score is defect probability multiplied by 0.5 plus 0.5 times the complexity index. No unavailable business criticality or historical test-cost variables were invented.")
    document.add_heading("Chapter 6: Dashboard and System Workflow", 1)
    document.add_paragraph("The Streamlit dashboard loads the persisted model bundle, displays dataset and model results, accepts a CSV of PC1 metrics, produces individual or batch predictions, and presents a ranked testing-priority table. The workflow is: raw dataset, preprocessing pipeline, model comparison, persisted model, probability risk, and complexity-aware prioritization.")
    document.add_heading("Chapter 7: Limitations and Future Scope", 1)
    document.add_paragraph("PC1 represents one NASA software context and may not generalize to other projects. Reported defects may not represent every defect, and the dataset lacks business priority, test execution cost, and complete test-suite history. Future work can evaluate additional projects, probability calibration, explainable AI, CI/CD integration, continuous retraining, and automated test generation.")
    document.add_heading("References", 1)
    document.add_paragraph("Sayyad Shirabad, J. and Menzies, T. J. (2005). The PROMISE Repository of Software Engineering Databases. School of Information Technology and Engineering, University of Ottawa.")
    document.add_paragraph("Shepperd, M., Song, Q., Sun, Z., and Mair, C. (2013). Data Quality: Some Comments on the NASA Software Defect Datasets. IEEE Transactions on Software Engineering, 39.")
    document.add_paragraph("Pedregosa, F. et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825-2830.")
    document.add_heading("Appendix: Preprocessing Record", 1)
    for key, value in summary.items():
        document.add_paragraph(f"{key}: {value}")
    document.save(reports / "final_report.docx")
    print("Created", reports / "final_report.docx")


if __name__ == "__main__":
    main()
